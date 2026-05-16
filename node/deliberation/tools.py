"""
OCC tool registry — explicit user-triggered tools only.
Tools are never called autonomously; they require explicit user intent.
File operations are sandboxed to the workspace/ directory inside the OCC repo.
"""

import os
import re
from pathlib import Path

_WORKSPACE: Path | None = None
_UPLOAD: Path | None = None
_PACKS_ROOT: Path | None = None


def set_workspace(path: Path):
    global _WORKSPACE
    _WORKSPACE = path


def set_upload(path: Path):
    global _UPLOAD
    _UPLOAD = path


def set_packs_root(path: Path):
    """Wire the expert-packs/ root so `download_pack` knows where to save."""
    global _PACKS_ROOT
    _PACKS_ROOT = path


def _safe_path(filename: str) -> Path:
    if _WORKSPACE is None:
        raise RuntimeError("Workspace not initialized.")
    resolved = (_WORKSPACE / filename).resolve()
    if not str(resolved).startswith(str(_WORKSPACE.resolve())):
        raise ValueError(f"Path '{filename}' escapes the workspace.")
    return resolved


def _safe_upload_path(filename: str) -> Path:
    if _UPLOAD is None:
        raise RuntimeError("Upload folder not initialized.")
    resolved = (_UPLOAD / filename).resolve()
    if not str(resolved).startswith(str(_UPLOAD.resolve())):
        raise ValueError(f"Path '{filename}' escapes the upload folder.")
    return resolved


def list_upload_files() -> list[str]:
    """Return filenames currently in the upload folder, excluding dotfiles
    and .gitkeep. Used by skills that need to enumerate user-uploaded files."""
    if _UPLOAD is None or not _UPLOAD.exists():
        return []
    return sorted(
        f.name for f in _UPLOAD.iterdir()
        if f.is_file() and not f.name.startswith(".")
    )

_URL_RE = re.compile(r'https?://\S+')

# Language-agnostic single keywords
_TOOL_KEYWORDS = {"web", "online", "internet", "workspace", "url", "file"}

# Phrases that clearly signal tool intent
_TOOL_PHRASES = {
    # web search
    "search the web", "search online", "look up online",
    "find online", "web search", "browse the web",
    "latest news", "current news", "news about",
    # file operations
    "write file", "write a file", "create file", "create a file",
    "save file", "save to file", "save as",
    "read file", "open file", "load file",
    "list files", "show files", "list the files",
    "in workspace", "to workspace", "into workspace",
    # code execution
    "run code", "execute code", "run this code", "run the code",
    "run script", "execute script",
}


def is_tool_request(query: str) -> bool:
    """True if the query explicitly requests a tool (web, file, or code execution)."""
    q = query.lower()
    if _URL_RE.search(query):
        return True
    if any(p in q for p in _TOOL_PHRASES):
        return True
    words = set(re.findall(r'\b\w+\b', q))
    return bool(words & _TOOL_KEYWORDS)


# ── Web-tool gating ────────────────────────────────────────────────────────────
# OCC's design rule: knowledge must come from community-approved expert packs.
# Web tools (web_search, fetch_url) bypass that guarantee, so they are exposed
# to the model ONLY when the user explicitly invokes the web.

WEB_TOOL_NAMES = frozenset({"web_search", "fetch_url"})

_WEB_KEYWORDS = {"web", "online", "internet"}
_WEB_PHRASES = {
    "search the web", "search online", "look up online", "find online",
    "web search", "browse the web", "browse online",
    "latest news", "current news", "news about",
    "cerca sul web", "cerca online", "cerca su internet",
    "su internet", "sul web",
}


def is_web_request(query: str) -> bool:
    """True iff the user explicitly invokes the web. URL in the message counts as explicit consent."""
    q = (query or "").lower()
    if _URL_RE.search(query or ""):
        return True
    if any(p in q for p in _WEB_PHRASES):
        return True
    words = set(re.findall(r'\b\w+\b', q))
    return bool(words & _WEB_KEYWORDS)


def get_allowed_tools(query: str) -> tuple[list, frozenset]:
    """
    Return (schema, allowed_function_names) for the given user query.
    Web tools (web_search, fetch_url) are stripped unless the query explicitly
    invokes the web. Other tools are always available.
    """
    if is_web_request(query):
        return TOOL_SCHEMA, frozenset(TOOL_FUNCTIONS.keys())
    schema = [t for t in TOOL_SCHEMA if t["function"]["name"] not in WEB_TOOL_NAMES]
    allowed = frozenset(TOOL_FUNCTIONS.keys()) - WEB_TOOL_NAMES
    return schema, allowed


def web_search(query: str, max_results: int = 5) -> str:
    from ddgs import DDGS
    results = DDGS().text(query, max_results=max_results)
    if not results:
        return "No results found."
    lines = []
    for r in results:
        lines.append(f"**{r['title']}**\n{r['body']}\nSource: {r['href']}\n")
    return "\n".join(lines)


def _url_is_safe(url: str) -> tuple[bool, str]:
    """Resolve the URL's host and reject internal targets.

    Defence against SSRF: a prompt-injected pack page or attacker-supplied
    URL could otherwise reach back into the user's own machine (Node GUI on
    127.0.0.1:7891, admin Hub on :7892), router admin pages (192.168.x),
    cloud metadata services (169.254.169.254), or VPN-internal services.
    Only http/https schemes are accepted. Returns (True, "") when safe.
    """
    from urllib.parse import urlparse
    import socket
    import ipaddress

    try:
        parsed = urlparse(url)
    except Exception as e:
        return False, f"invalid URL: {e}"
    if parsed.scheme not in ("http", "https"):
        return False, f"scheme '{parsed.scheme}' not allowed (use http/https)"
    host = (parsed.hostname or "").strip()
    if not host:
        return False, "missing host"
    try:
        infos = socket.getaddrinfo(host, None)
    except Exception as e:
        return False, f"cannot resolve host: {e}"
    for info in infos:
        addr = info[4][0]
        try:
            ip = ipaddress.ip_address(addr)
        except ValueError:
            continue
        if (ip.is_private or ip.is_loopback or ip.is_link_local
                or ip.is_multicast or ip.is_reserved or ip.is_unspecified):
            return False, f"host resolves to non-public address {addr}"
    return True, ""


def fetch_url(url: str) -> str:
    """Fetch a URL and return its main text content.

    Two-stage extraction:
      1. trafilatura (best-quality main-content extraction).
      2. If trafilatura returns nothing (JS-heavy SPA, paywall stub, etc.),
         fall back to a crude HTML strip so we at least return *some* text
         instead of an empty result.
    """
    import re
    import httpx
    import trafilatura
    ok, why = _url_is_safe(url)
    if not ok:
        return f"Refused to fetch {url}: {why}."
    try:
        with httpx.Client(
            timeout=15.0,
            follow_redirects=False,
            headers={
                "User-Agent": "OCC-Node/1.0 (+https://github.com/VickyWenSZ/occ)",
                "Accept-Language": "it,en;q=0.9",
            },
        ) as client:
            resp = client.get(url)
            # Manual redirect handling: re-validate the target each hop so a
            # public host can't redirect us to an internal one.
            hops = 0
            while resp.is_redirect and hops < 5:
                next_url = resp.headers.get("location", "")
                if not next_url:
                    break
                from urllib.parse import urljoin
                next_url = urljoin(str(resp.url), next_url)
                ok, why = _url_is_safe(next_url)
                if not ok:
                    return f"Refused to follow redirect to {next_url}: {why}."
                resp = client.get(next_url)
                hops += 1
            resp.raise_for_status()
    except httpx.HTTPStatusError as e:
        return f"HTTP {e.response.status_code} on {url}"
    except httpx.TimeoutException:
        return f"Timeout fetching {url} (15s)"
    except Exception as e:
        return f"Could not fetch URL ({type(e).__name__}): {e}"

    html = resp.text
    text = trafilatura.extract(html, include_comments=False, include_tables=False) or ""
    if text.strip():
        return text[:8000] if len(text) > 8000 else text

    # Fallback: trafilatura returned empty (paywall, SPA, anti-bot stub).
    # Strip script/style blocks then strip remaining tags.
    cleaned = re.sub(r"<script[^>]*>.*?</script>", " ", html, flags=re.DOTALL | re.IGNORECASE)
    cleaned = re.sub(r"<style[^>]*>.*?</style>", " ", cleaned, flags=re.DOTALL | re.IGNORECASE)
    cleaned = re.sub(r"<[^>]+>", " ", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    if not cleaned:
        return f"Could not extract readable text from {url} (page returned empty after both extractors)."
    # Tag the fallback so Qwen knows this is rough
    return ("[fallback extraction — main-content extractor returned empty]\n"
            + (cleaned[:8000] if len(cleaned) > 8000 else cleaned))


def read_file(filename: str) -> str:
    try:
        path = _safe_upload_path(filename)
        if not path.exists():
            return f"File '{filename}' not found in upload folder."
        return path.read_text(encoding="utf-8")
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error reading file: {e}"


def write_file(filename: str, content: str) -> str:
    try:
        path = _safe_path(filename)
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(content, encoding="utf-8")
        return f"File '{filename}' written successfully."
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error writing file: {e}"


def list_files(subdir: str = "") -> str:
    try:
        base = _safe_path(subdir) if subdir else _WORKSPACE
        if not base.exists():
            return f"Directory '{subdir}' not found in workspace."
        entries = sorted(base.iterdir(), key=lambda p: (p.is_file(), p.name))
        if not entries:
            return "Workspace is empty."
        lines = []
        for p in entries:
            if p.name == ".gitkeep":
                continue
            marker = "/" if p.is_dir() else ""
            lines.append(f"  {p.name}{marker}")
        return "Workspace contents:\n" + "\n".join(lines)
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error listing files: {e}"


# Prelude injected before user code in run_code. Installs a Python audit hook
# that blocks the highest-impact escape vectors:
#   - subprocess / os.system / os.exec*  → no local privilege escalation
#   - socket networking + DNS            → no data exfiltration outbound
#   - ctypes                             → can't FFI around the audit hook
# Audit hooks (PEP 578) fire before the action and cannot be uninstalled from
# Python code. We pair this with a stripped env and POSIX rlimits in the
# caller. This is defence-in-depth against prompt-injected pack content that
# tries to coerce the model into calling run_code with malicious payloads —
# not a hard isolation boundary (use a container or VM for that).
_RUN_CODE_SANDBOX_PRELUDE = r'''
import sys as _occ_sys
_OCC_FORBIDDEN_EVENTS = frozenset({
    "subprocess.Popen",
    "os.system",
    "os.exec",
    "os.spawn",
    "socket.connect",
    "socket.bind",
    "socket.getaddrinfo",
    "socket.gethostbyname",
    "winreg.OpenKey",
    "winreg.CreateKey",
    "winreg.DeleteKey",
    "ctypes.dlopen",
    "ctypes.dlsym",
    "ctypes.cdata/buffer",
})
_OCC_FORBIDDEN_MODULES = frozenset({
    "ctypes", "_ctypes", "ctypes.util",
})
def _occ_audit(event, args):
    if event in _OCC_FORBIDDEN_EVENTS:
        raise PermissionError("sandbox: '" + event + "' is not allowed in run_code")
    if event == "import":
        mod = args[0] if args else ""
        if mod in _OCC_FORBIDDEN_MODULES:
            raise PermissionError("sandbox: importing '" + mod + "' is not allowed in run_code")
_occ_sys.addaudithook(_occ_audit)
del _occ_sys, _OCC_FORBIDDEN_EVENTS, _OCC_FORBIDDEN_MODULES, _occ_audit
'''

# Cap stdout/stderr returned to the model. Without this, an infinite-print
# loop in 30s can produce hundreds of MB and explode the next LLM call's ctx.
_RUN_CODE_OUTPUT_CAP = 8000


def run_code(code: str) -> str:
    import subprocess
    import sys

    # Minimal environment. Inherit only what the child Python actually needs
    # to import its stdlib (SystemRoot on Windows, PATH stripped). No HOME →
    # the user's dotfiles aren't directly addressable via ~. No PYTHONPATH so
    # the user can't load arbitrary modules from outside.
    env = {}
    if os.name == "nt":
        for k in ("SYSTEMROOT", "SystemRoot", "WINDIR", "PATHEXT", "TEMP", "TMP"):
            if k in os.environ:
                env[k] = os.environ[k]
    env["PYTHONDONTWRITEBYTECODE"] = "1"
    env["PYTHONUTF8"] = "1"

    # POSIX-only: cap CPU time, address space, and core size. Set via
    # preexec_fn in the child before exec, so the limits apply to user code.
    preexec = None
    if os.name == "posix":
        def _set_limits():
            import resource
            # Memory cap: 512 MB virtual address space. Generous for legit
            # computation, blocks fork-bomb-style ballooning.
            resource.setrlimit(resource.RLIMIT_AS, (512 * 1024 * 1024, 512 * 1024 * 1024))
            resource.setrlimit(resource.RLIMIT_CPU, (35, 35))   # ~30s wall + headroom
            resource.setrlimit(resource.RLIMIT_CORE, (0, 0))
            resource.setrlimit(resource.RLIMIT_NPROC, (32, 32))
        preexec = _set_limits

    full_code = _RUN_CODE_SANDBOX_PRELUDE + "\n" + code
    try:
        result = subprocess.run(
            [sys.executable, "-I", "-c", full_code],
            capture_output=True,
            text=True,
            timeout=30,
            cwd=str(_WORKSPACE),
            env=env,
            preexec_fn=preexec,
        )
        out = (result.stdout or "")[:_RUN_CODE_OUTPUT_CAP]
        err = (result.stderr or "")[:_RUN_CODE_OUTPUT_CAP]
        output = out
        if err:
            output += f"\n[stderr]\n{err}"
        if (len(result.stdout or "") > _RUN_CODE_OUTPUT_CAP
                or len(result.stderr or "") > _RUN_CODE_OUTPUT_CAP):
            output += "\n[...output truncated]"
        return output.strip() or "(no output)"
    except subprocess.TimeoutExpired:
        return "Error: code execution timed out (30s limit)."
    except Exception as e:
        return f"Error running code: {e}"


# ── Document readers (PDF / DOCX / XLSX) ───────────────────────────────────────

_DOC_READ_LIMIT = 6000  # max chars returned to keep tool result inside chat ctx


def read_pdf(filename: str) -> str:
    try:
        path = _safe_upload_path(filename)
        if not path.exists():
            return f"File '{filename}' not found in upload folder."
        import fitz  # PyMuPDF
        doc = fitz.open(str(path))
        try:
            parts = []
            total = 0
            for page in doc:
                t = page.get_text()
                if not t:
                    continue
                parts.append(t)
                total += len(t)
                if total >= _DOC_READ_LIMIT:
                    break
            text = "\n\n".join(parts)
        finally:
            doc.close()
        if not text.strip():
            return f"PDF '{filename}' has no extractable text (possibly scanned)."
        if len(text) > _DOC_READ_LIMIT:
            text = text[:_DOC_READ_LIMIT] + "\n\n[...truncated]"
        return text
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error reading PDF: {e}"


def read_docx(filename: str) -> str:
    try:
        path = _safe_upload_path(filename)
        if not path.exists():
            return f"File '{filename}' not found in upload folder."
        from docx import Document
        doc = Document(str(path))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                parts.append(" | ".join(cells))
        text = "\n".join(parts)
        if not text.strip():
            return f"Document '{filename}' is empty."
        if len(text) > _DOC_READ_LIMIT:
            text = text[:_DOC_READ_LIMIT] + "\n\n[...truncated]"
        return text
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error reading docx: {e}"


def read_xlsx(filename: str) -> str:
    try:
        path = _safe_upload_path(filename)
        if not path.exists():
            return f"File '{filename}' not found in upload folder."
        from openpyxl import load_workbook
        wb = load_workbook(str(path), data_only=True, read_only=True)
        try:
            parts = []
            running = 0
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                header = f"--- Sheet: {sheet_name} ---"
                parts.append(header)
                running += len(header)
                for row in sheet.iter_rows(values_only=True):
                    row_str = "\t".join("" if v is None else str(v) for v in row)
                    if not row_str.strip():
                        continue
                    parts.append(row_str)
                    running += len(row_str)
                    if running >= _DOC_READ_LIMIT:
                        break
                if running >= _DOC_READ_LIMIT:
                    break
        finally:
            wb.close()
        text = "\n".join(parts)
        if not text.strip():
            return f"Workbook '{filename}' is empty."
        if len(text) > _DOC_READ_LIMIT:
            text = text[:_DOC_READ_LIMIT] + "\n\n[...truncated]"
        return text
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error reading xlsx: {e}"


# ── Audio transcription (Whisper) ──────────────────────────────────────────────

_WHISPER_MODEL = None


def _get_whisper_model():
    """Lazy-load the Whisper model; downloaded once on first use."""
    global _WHISPER_MODEL
    if _WHISPER_MODEL is None:
        import os
        from faster_whisper import WhisperModel
        model_size = os.environ.get("OCC_WHISPER_MODEL", "base")
        device = os.environ.get("OCC_WHISPER_DEVICE", "cpu")
        compute_type = "int8" if device == "cpu" else "float16"
        _WHISPER_MODEL = WhisperModel(model_size, device=device, compute_type=compute_type)
    return _WHISPER_MODEL


def transcribe_audio(filename: str) -> str:
    try:
        path = _safe_upload_path(filename)
        if not path.exists():
            return f"Audio file '{filename}' not found in upload folder."
        model = _get_whisper_model()
        segments, info = model.transcribe(str(path), beam_size=5)
        parts = [seg.text.strip() for seg in segments if seg.text and seg.text.strip()]
        if not parts:
            return f"No speech detected in '{filename}'."
        text = " ".join(parts)
        return f"[language: {info.language}, duration: {info.duration:.1f}s]\n\n{text}"
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error transcribing audio: {e}"


TOOL_SCHEMA = [
    {
        "type": "function",
        "function": {
            "name": "web_search",
            "description": "Search the web for current information. Use when the user asks to search online.",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {"type": "string", "description": "The search query"},
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "fetch_url",
            "description": "Fetch and read the content of a specific web page. Use when the user provides a URL or asks to read a page.",
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {"type": "string", "description": "The full URL to fetch"},
                },
                "required": ["url"],
            },
        },
    },
]

TOOL_SCHEMA += [
    {
        "type": "function",
        "function": {
            "name": "read_file",
            "description": "Read a file from the workspace directory.",
            "parameters": {
                "type": "object",
                "properties": {
                    "filename": {"type": "string", "description": "Filename or relative path within workspace"},
                },
                "required": ["filename"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "write_file",
            "description": "Write or create a file in the workspace directory.",
            "parameters": {
                "type": "object",
                "properties": {
                    "filename": {"type": "string", "description": "Filename or relative path within workspace"},
                    "content": {"type": "string", "description": "Content to write"},
                },
                "required": ["filename", "content"],
            },
        },
    },
]

TOOL_SCHEMA += [
    {
        "type": "function",
        "function": {
            "name": "list_files",
            "description": "List files and directories in the workspace. Use when the user asks what files exist or before reading a file.",
            "parameters": {
                "type": "object",
                "properties": {
                    "subdir": {"type": "string", "description": "Optional subdirectory within workspace (leave empty for root)"},
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "run_code",
            "description": "Execute a Python code snippet and return its output. The working directory is the workspace.",
            "parameters": {
                "type": "object",
                "properties": {
                    "code": {"type": "string", "description": "Python code to execute"},
                },
                "required": ["code"],
            },
        },
    },
]

TOOL_SCHEMA += [
    {
        "type": "function",
        "function": {
            "name": "read_pdf",
            "description": "Read text content from a PDF file in the upload folder. Use when the user asks to read, summarize, or answer questions about a PDF document they attached.",
            "parameters": {
                "type": "object",
                "properties": {
                    "filename": {"type": "string", "description": "Filename of the uploaded PDF (e.g. 'report.pdf')"},
                },
                "required": ["filename"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_docx",
            "description": "Read text content from a Microsoft Word .docx file in the upload folder.",
            "parameters": {
                "type": "object",
                "properties": {
                    "filename": {"type": "string", "description": "Filename of the uploaded .docx (e.g. 'document.docx')"},
                },
                "required": ["filename"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_xlsx",
            "description": "Read cell contents from a Microsoft Excel .xlsx file in the upload folder. Returns all sheets with rows as tab-separated values.",
            "parameters": {
                "type": "object",
                "properties": {
                    "filename": {"type": "string", "description": "Filename of the uploaded .xlsx (e.g. 'data.xlsx')"},
                },
                "required": ["filename"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "transcribe_audio",
            "description": "Transcribe an audio file in the upload folder to text using Whisper. Supports common formats (mp3, wav, m4a, ogg, flac).",
            "parameters": {
                "type": "object",
                "properties": {
                    "filename": {"type": "string", "description": "Filename of the uploaded audio (e.g. 'recording.mp3')"},
                },
                "required": ["filename"],
            },
        },
    },
]

# ── Broker pack catalog: list + download ─────────────────────────────────────
# Lets the user browse the OCC broker's pack tree from chat and pull packs
# locally. Packs are saved DISABLED — the Settings toggle controls activation,
# and the existing /api/local/pack-state endpoint triggers `reindex_pack` on
# enable so the FTS5 index is built lazily.

_PACK_PATH_RE = re.compile(r"^[a-z0-9][a-z0-9/_\-]*$")

# Per-call cache for the broker's bulk summaries dict. list_packs probes many
# children in a single invocation; one fetch per call is enough.
_PACK_SUMMARIES_CACHE: dict | None = None


def _broker_tree(prefix: str):
    """GET /tree[/prefix]. Returns the parsed JSON or None on any failure."""
    import httpx
    from node.retrieval.pack_cache import SERVER_URL
    url = f"{SERVER_URL}/tree" if not prefix else f"{SERVER_URL}/tree/{prefix}"
    try:
        with httpx.Client(timeout=8.0, follow_redirects=True) as client:
            resp = client.get(url)
        if resp.status_code != 200:
            return None
        return resp.json()
    except Exception:
        return None


def _fetch_summaries_bulk() -> dict:
    """GET /packs/summaries — one bulk call returns {pack_path: summary} for
    every indexed pack on the broker. Cheaper than N per-pack manifest calls."""
    import httpx
    from node.retrieval.pack_cache import SERVER_URL
    try:
        with httpx.Client(timeout=8.0, follow_redirects=True) as client:
            resp = client.get(f"{SERVER_URL}/packs/summaries")
        if resp.status_code != 200:
            return {}
        data = resp.json()
        summaries = data.get("summaries", {}) if isinstance(data, dict) else {}
        return summaries if isinstance(summaries, dict) else {}
    except Exception:
        return {}


def _fetch_pack_summary_remote(pack_path: str) -> str:
    """Return a single pack's summary from the cached bulk dict."""
    global _PACK_SUMMARIES_CACHE
    if _PACK_SUMMARIES_CACHE is None:
        _PACK_SUMMARIES_CACHE = _fetch_summaries_bulk()
    return str(_PACK_SUMMARIES_CACHE.get(pack_path, "") or "").strip()


def _pack_installed_locally(pack_path: str) -> bool:
    """True when expert-packs/<pack_path>/ has both manifest.yaml and wiki/index.md."""
    if _PACKS_ROOT is None:
        return False
    p = _PACKS_ROOT / pack_path
    return (p / "manifest.yaml").exists() and (p / "wiki" / "index.md").exists()


def list_packs(prefix: str = "") -> str:
    """Browse the OCC broker's pack catalog at the given prefix.

    Returns a human-readable listing of immediate children at that path so the
    user can navigate the tree conversationally. Pack leaves are annotated
    with their manifest summary and a marker when already installed locally.
    """
    global _PACK_SUMMARIES_CACHE
    _PACK_SUMMARIES_CACHE = None  # refresh once per call so re-browses see updates
    prefix_clean = (prefix or "").strip().strip("/").replace("\\", "/")
    if prefix_clean and not _PACK_PATH_RE.match(prefix_clean):
        return f"Invalid catalog path: '{prefix}'."

    data = _broker_tree(prefix_clean)
    if data is None:
        return "Could not reach the OCC broker."

    if not prefix_clean:
        children = data if isinstance(data, list) else []
        is_pack_here = False
    else:
        if not isinstance(data, dict):
            return f"Nothing at '{prefix_clean}' on the broker."
        children = data.get("children", []) or []
        is_pack_here = bool(data.get("has_pack"))

    lines: list[str] = []
    location = prefix_clean or "(root)"

    if is_pack_here:
        summary = _fetch_pack_summary_remote(prefix_clean)
        marker = " — installed locally" if _pack_installed_locally(prefix_clean) else ""
        lines.append(f"Pack: {prefix_clean}{marker}")
        if summary:
            lines.append(f"  {summary}")
        else:
            lines.append("  (no description available — pack manifest has empty summary)")
        lines.append("")

    if children:
        lines.append(f"Under {location}:")
        all_summaries = _PACK_SUMMARIES_CACHE or {}
        for name in children:
            if not isinstance(name, str) or not name:
                continue
            child_path = f"{prefix_clean}/{name}" if prefix_clean else name
            child_data = _broker_tree(child_path)
            child_has_pack = bool(
                isinstance(child_data, dict) and child_data.get("has_pack")
            )
            if child_has_pack:
                summary = _fetch_pack_summary_remote(child_path)
                marker = " ✓" if _pack_installed_locally(child_path) else ""
                line = f"  - {name}{marker} (pack)"
                if summary:
                    snippet = summary[:140].rstrip()
                    if len(summary) > 140:
                        snippet += "..."
                    line += f" — {snippet}"
                else:
                    line += " — (no description available yet)"
                lines.append(line)
            else:
                # Count packs reachable under this category so the user knows
                # whether it's worth drilling down further.
                prefix_for_count = f"{child_path}/"
                pack_count = sum(
                    1 for pp in all_summaries if pp.startswith(prefix_for_count)
                )
                if pack_count == 1:
                    lines.append(f"  - {name}/ (category, 1 pack inside)")
                elif pack_count > 1:
                    lines.append(f"  - {name}/ (category, {pack_count} packs inside)")
                else:
                    lines.append(f"  - {name}/ (category)")
    elif not is_pack_here:
        return f"Nothing at '{prefix_clean}' on the broker."

    return "\n".join(lines).strip()


def download_pack(path: str, force: bool = False) -> str:
    """Download a pack from the OCC broker to the local expert-packs/ folder.

    Saves the pack as DISABLED so it doesn't interfere with current retrieval.
    The user activates it from Settings — that toggle calls reindex_pack so
    the pack becomes searchable on first activation.

    Two-step confirmation when a pack of the same path already exists locally:
    the first call returns a message describing the conflict and does NOT
    touch the filesystem. The LLM is expected to relay that to the user and,
    on confirmation, re-call this tool with force=true.
    """
    import httpx
    from node.retrieval.pack_cache import SERVER_URL

    pack_path = (path or "").strip().strip("/").replace("\\", "/")
    if not pack_path:
        return "download_pack: missing pack path."
    if not _PACK_PATH_RE.match(pack_path):
        return f"download_pack: invalid pack path '{path}'."
    if _PACKS_ROOT is None:
        return "download_pack: local packs root not configured."

    pack_dir = _PACKS_ROOT / pack_path
    wiki_dir = pack_dir / "wiki"
    pre_existed = pack_dir.exists() and (pack_dir / "manifest.yaml").exists()

    if pre_existed and not force:
        # Surface the conflict to the user via the model. Count what's at risk
        # so the user can make an informed call.
        try:
            existing_pages = sum(
                1 for f in (wiki_dir).rglob("*.md") if f.is_file()
            ) if wiki_dir.exists() else 0
        except Exception:
            existing_pages = 0
        try:
            from node.apps.cli.config import load_disabled_packs
            enabled = pack_path not in set(load_disabled_packs())
        except Exception:
            enabled = False
        status = "enabled" if enabled else "disabled"
        return (
            f"A local pack already exists at '{pack_path}' "
            f"({existing_pages} pages, currently {status}). "
            f"Re-downloading will OVERWRITE every file on disk, including any "
            f"local edits made with Forge. "
            f"If the user confirms they want to overwrite, call this tool "
            f"again with force=true. If they want to keep the local copy, "
            f"do nothing."
        )

    pages_to_fetch: list[str] = []
    try:
        with httpx.Client(timeout=20.0, follow_redirects=True) as client:
            # index.md — the source of truth for which pages belong to the pack
            r = client.get(f"{SERVER_URL}/packs/{pack_path}/wiki/index.md")
            if r.status_code != 200:
                return (
                    f"download_pack: pack '{pack_path}' not found on broker "
                    f"(HTTP {r.status_code})."
                )
            index_md = r.text
            for line in index_md.splitlines():
                if not line.startswith("|"):
                    continue
                m = re.match(r"^\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|", line)
                if not m:
                    continue
                page_file = m.group(1).strip()
                if page_file.lower() == "file" or page_file.startswith("-"):
                    continue
                if page_file:
                    pages_to_fetch.append(page_file)

            wiki_dir.mkdir(parents=True, exist_ok=True)
            (wiki_dir / "index.md").write_text(index_md, encoding="utf-8")

            pages_ok = 0
            pages_failed: list[str] = []
            # Allow nested page paths (e.g. "concepts/julius-caesar.md") but
            # block traversal and absolute paths. Each segment is restricted
            # to safe characters.
            page_re = re.compile(r"^[A-Za-z0-9._\-]+(/[A-Za-z0-9._\-]+)*\.md$")
            for page_file in pages_to_fetch:
                if not page_re.match(page_file) or ".." in page_file:
                    pages_failed.append(page_file)
                    continue
                try:
                    pr = client.get(
                        f"{SERVER_URL}/packs/{pack_path}/wiki/{page_file}"
                    )
                    if pr.status_code != 200:
                        pages_failed.append(page_file)
                        continue
                    out_path = wiki_dir / page_file
                    out_path.parent.mkdir(parents=True, exist_ok=True)
                    out_path.write_text(pr.text, encoding="utf-8")
                    pages_ok += 1
                except Exception:
                    pages_failed.append(page_file)

            mr = client.get(f"{SERVER_URL}/packs/{pack_path}/manifest.yaml")
            if mr.status_code != 200:
                return (
                    f"download_pack: pack '{pack_path}' missing manifest.yaml "
                    f"on broker (HTTP {mr.status_code})."
                )
            (pack_dir / "manifest.yaml").write_text(mr.text, encoding="utf-8")
    except Exception as e:
        return f"download_pack: fetch failed: {type(e).__name__}: {e}"

    # Stamp the new pack as disabled. Existing entries are preserved.
    try:
        from node.apps.cli.config import load_disabled_packs, save_disabled_packs
        disabled = set(load_disabled_packs())
        disabled.add(pack_path)
        save_disabled_packs(list(disabled))
    except Exception:
        pass

    notes = []
    if pre_existed:
        notes.append("overwrote existing local copy")
    if pages_failed:
        notes.append(f"{len(pages_failed)} pages failed: {', '.join(pages_failed[:3])}")
    suffix = f" ({'; '.join(notes)})" if notes else ""

    return (
        f"Pack '{pack_path}' downloaded: {pages_ok} pages + manifest{suffix}.\n"
        f"Saved as DISABLED. Activate it from Settings to use it in local mode — "
        f"activation triggers the FTS5 index build automatically."
    )


TOOL_SCHEMA += [
    {
        "type": "function",
        "function": {
            "name": "list_packs",
            "description": (
                "Browse the OCC broker's catalog of knowledge packs available "
                "for remote use or local download. Returns one level of the "
                "catalog tree at the given prefix (empty prefix = top level). "
                "Use when the user asks what packs are available on the broker, "
                "wants to browse the catalog, or wants to check if a topic is "
                "covered."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "prefix": {
                        "type": "string",
                        "description": (
                            "Catalog path to browse. Empty string ('') for the "
                            "top-level categories. Use a path like 'marketing' "
                            "or 'history/ancient-rome' to drill down."
                        ),
                    },
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "download_pack",
            "description": (
                "Download a pack from the OCC broker to the local expert-packs/ "
                "folder so the user can use it later in local mode or modify it "
                "with Forge. The pack is saved DISABLED — activation from "
                "Settings makes it searchable (and triggers indexing). Use when "
                "the user explicitly asks to download or save a pack. "
                "If a pack with the same path already exists locally, the first "
                "call returns a conflict message asking for confirmation without "
                "touching the filesystem; relay that to the user and only re-call "
                "with force=true if they explicitly confirm overwriting."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": (
                            "Full pack path on the broker (e.g. 'marketing/"
                            "storybrand' or 'history/ancient-rome/caesar')."
                        ),
                    },
                    "force": {
                        "type": "boolean",
                        "description": (
                            "Set to true only after the user has confirmed they "
                            "want to overwrite a pre-existing local copy of the "
                            "same pack. Default false."
                        ),
                    },
                },
                "required": ["path"],
            },
        },
    },
]


TOOL_FUNCTIONS = {
    "web_search": web_search,
    "fetch_url": fetch_url,
    "read_file": read_file,
    "write_file": write_file,
    "list_files": list_files,
    "run_code": run_code,
    "read_pdf": read_pdf,
    "read_docx": read_docx,
    "read_xlsx": read_xlsx,
    "transcribe_audio": transcribe_audio,
    "list_packs": list_packs,
    "download_pack": download_pack,
}
